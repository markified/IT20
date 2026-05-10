"""Rebuild DOCUMENTATION IT20.docx in IEEE-style two-column format.

- Single-column section: Title, author, Abstract, Keywords.
- Two-column continuous section: Body (1. Introduction ... 3.2 Conclusion).
- Concise paragraphs, figures inline for justification (IT14-style).
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

ROOT = Path(r"C:\Users\ADMIN\Desktop\IT20")
IMG = ROOT / "extracted_imgs"
EDA = ROOT / "eda_imgs"
OUT = ROOT / "DOCUMENTATION IT20.docx"

doc = Document()

# Default font
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(10)

# ----- Section 0: single column (title block + abstract) -----
sec0 = doc.sections[0]
sec0.page_height = Cm(29.7)
sec0.page_width = Cm(21.0)
sec0.top_margin = Cm(1.9)
sec0.bottom_margin = Cm(1.9)
sec0.left_margin = Cm(1.5)
sec0.right_margin = Cm(1.5)


def set_columns(section, num=2, sep=True):
    sectPr = section._sectPr
    cols = sectPr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sectPr.append(cols)
    cols.set(qn("w:num"), str(num))
    cols.set(qn("w:equalWidth"), "1")
    cols.set(qn("w:space"), "432")  # ~0.3 inch gutter
    if sep and num > 1:
        cols.set(qn("w:sep"), "1")


def add_para(text, *, bold=False, italic=False, size=10, align=None,
             space_after=2, space_before=0, font="Times New Roman"):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    r = p.add_run(text)
    r.font.name = font
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    return p


def add_heading(text, *, level=1):
    sizes = {1: 11, 2: 10, 3: 10}
    bolds = {1: True, 2: True, 3: True}
    italics = {1: False, 2: False, 3: True}
    return add_para(text, bold=bolds[level], italic=italics[level],
                    size=sizes[level], space_before=6, space_after=2)


def add_figure(path: Path, caption: str, width_inches=3.2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run()
    r.add_picture(str(path), width=Inches(width_inches))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(6)
    cr = cap.add_run(caption)
    cr.italic = True
    cr.font.size = Pt(9)
    cr.font.name = "Times New Roman"


# ===== TITLE BLOCK (single column) =====
add_para(
    "Predictive Analytics for Early Identification of Disputed "
    "Accounts Receivable Invoices",
    bold=True, size=18, align=WD_ALIGN_PARAGRAPH.CENTER,
    space_before=6, space_after=8,
)
add_para("Mark Clarence P. Balmes",
         size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para("College of Computing Education, University of Mindanao",
         size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para("Matina, Davao City, Philippines",
         size=10, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para("mbalmes@umindanao.edu.ph",
         size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)

# Abstract
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(6)
r = p.add_run("Abstract — ")
r.bold = True
r.italic = True
r.font.size = Pt(10)
r.font.name = "Times New Roman"
r2 = p.add_run(
    "Disputed customer invoices delay cash collection and inflate "
    "operating cost in finance shared-service centers. This study "
    "applies predictive analytics on the IBM Accounts Receivable "
    "dataset (2,466 invoices, 22.7% disputed) to flag at-risk "
    "invoices at issuance time. Two leakage-free classifiers — "
    "Logistic Regression (interpretable baseline) and Random Forest "
    "(ensemble) — were trained on six engineered features and refined "
    "through SelectKBest feature selection, GridSearchCV "
    "hyper-parameter tuning, and probability-threshold calibration on "
    "the precision–recall curve. The refined Random Forest delivered "
    "the best results (Accuracy 0.7713, Precision 0.4946, Recall 0.4107, "
    "F1 0.4488, ROC-AUC 0.7460), a +33% gain in precision and +20% in "
    "ROC-AUC over its baseline, while keeping a workable recall for "
    "early-warning use. The model is recommended for production as a "
    "ranked-risk queue for the collections team."
)
r2.italic = True
r2.font.size = Pt(10)
r2.font.name = "Times New Roman"

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(8)
r = p.add_run("Keywords — ")
r.bold = True
r.italic = True
r.font.size = Pt(10)
r.font.name = "Times New Roman"
r2 = p.add_run(
    "predictive analytics, accounts receivable, dispute prediction, "
    "logistic regression, random forest, hyperparameter tuning, "
    "threshold calibration"
)
r2.italic = True
r2.font.size = Pt(10)
r2.font.name = "Times New Roman"

# ----- Insert continuous section break, switch to 2 columns -----
new_section = doc.add_section(WD_SECTION.CONTINUOUS)
new_section.top_margin = Cm(1.9)
new_section.bottom_margin = Cm(1.9)
new_section.left_margin = Cm(1.5)
new_section.right_margin = Cm(1.5)
set_columns(new_section, num=2, sep=False)

# ===== 1. INTRODUCTION =====
add_heading("1. Introduction", level=1)

add_heading("1.1 Business Problem Description", level=2)
add_para(
    "Trade-credit invoices are a working-capital lever for B2B firms, "
    "but disputed invoices stall cash collection, raise Days Sales "
    "Outstanding, and consume collector time. The shared-service "
    "center wants to know — at the moment an invoice is issued — "
    "which customers are likely to dispute, so collectors can "
    "intervene proactively instead of reactively chasing aged "
    "receivables. Manual rules (country, amount thresholds) miss "
    "the joint patterns between customer behavior, invoice timing, "
    "and channel choice, which is why a predictive model is needed."
)

add_heading("1.2 Objectives", level=2)
add_para("1.2.1 To define the dispute-prediction problem and its "
         "significance to the AR collections workflow.")
add_para("1.2.2 To analyze the IBM Accounts Receivable dataset and "
         "engineer leakage-free features available at invoice "
         "issuance.")
add_para("1.2.3 To frame the task as binary classification and "
         "justify the metric choice (precision-oriented for a "
         "ranked-risk queue).")
add_para("1.2.4 To compare a human-implemented Logistic Regression "
         "with an AI-recommended Random Forest, then refine and "
         "select the final model on performance, stability, and "
         "business fit.")

# ===== 2. METHODOLOGY =====
add_heading("2. Methodology", level=1)

add_heading("2.1 Dataset Overview", level=2)
add_para(
    "Source: IBM Sample — WA_Fn-UseC_-Accounts-Receivable.csv. "
    "Size: 2,466 invoices × 12 attributes. "
    "Target: Disputed (Yes/No), 22.7% positive class — moderately "
    "imbalanced. Key fields: countryCode, customerID, PaperlessDate, "
    "InvoiceNumber, InvoiceDate, DueDate, InvoiceAmount, Disputed, "
    "PaperlessBill, DaysToSettle, DaysLate, SettledDate."
)

add_heading("2.2 Data Analysis", level=2)
add_para(
    "The raw file has zero missing values and no duplicates. Class "
    "distribution: 1,907 No / 559 Yes (22.7% disputed). InvoiceAmount "
    "is right-skewed; CreditTermDays (DueDate − InvoiceDate) is "
    "discrete and clusters at 30/45/60. Settlement-time fields "
    "(DaysToSettle, DaysLate, SettledDate) leak the future and were "
    "excluded from training."
)

add_figure(EDA / "eda_class_missing.png",
           "Figure A. Target class distribution and missing-value audit "
           "(0 missing across all columns).",
           width_inches=3.2)
add_para(
    "Proof: 22.7% positive class confirms moderate imbalance and "
    "justifies class_weight='balanced' and precision-oriented "
    "evaluation. The right panel verifies a clean dataset — no "
    "imputation required at the raw stage."
)

add_figure(EDA / "eda_invoice_amount.png",
           "Figure B. InvoiceAmount distribution and dispute split.",
           width_inches=3.2)
add_para(
    "Proof: InvoiceAmount is right-skewed and disputed invoices skew "
    "toward larger amounts — strong signal for the model and "
    "motivation to keep this feature."
)

add_figure(EDA / "eda_corr.png",
           "Figure C. Correlation matrix on raw numeric fields.",
           width_inches=3.0)
add_para(
    "Proof: low pair-wise correlation between predictor candidates — "
    "no severe multicollinearity, so all candidate features can enter "
    "the model."
)

add_heading("2.3 Data Preparation and Feature Engineering", level=2)
add_para(
    "Six leakage-free features were derived from issuance-time data: "
    "(1) CreditTermDays, (2) InvoiceMonth, (3) InvoiceDayOfWeek, "
    "(4) InvoiceQuarter, (5) DaysSincePaperless, "
    "(6) CustomerFreq (count of past invoices per customer). "
    "PaperlessBill was label-encoded. Numerical features were "
    "standardized for Logistic Regression only. The dataset was "
    "split 80/20 with stratification on Disputed."
)

add_figure(EDA / "fe_engineered_features.png",
           "Figure D. Distributions of engineered features — "
           "CreditTermDays, DaysSincePaperless, CustomerFreq.",
           width_inches=3.3)
add_para(
    "Proof of feature creation: CreditTermDays clusters at the "
    "expected 30/45/60-day terms; DaysSincePaperless and CustomerFreq "
    "show meaningful spread across customers — the engineered fields "
    "are well-defined and non-degenerate."
)

add_figure(EDA / "fe_target_relation.png",
           "Figure E. Engineered features versus dispute outcome.",
           width_inches=3.3)
add_para(
    "Proof of discriminative power: disputed invoices have visibly "
    "different CreditTermDays and CustomerFreq profiles, and the "
    "monthly dispute rate varies by ~10 pp — the new features carry "
    "real signal, not just noise."
)

add_figure(EDA / "fe_corr_refined.png",
           "Figure F. Correlation heat-map of the full refined feature "
           "matrix (with target).",
           width_inches=3.2)
add_para(
    "Proof: refined features remain weakly correlated with each "
    "other while several show non-trivial correlation with Disputed — "
    "supporting their inclusion and the use of SelectKBest to keep "
    "the most informative subset."
)

add_heading("2.7 Final Model Selection", level=2)
add_para(
    "The Refined Random Forest is selected as the final production "
    "model based on three criteria required by the business case:"
)
add_para(
    "(1) Performance — it achieves the highest score on every key "
    "metric: Accuracy 0.7713, Precision 0.4946, F1 0.4488, and "
    "ROC-AUC 0.7460. It is the only model that crosses the "
    "ROC-AUC 0.70 threshold considered usable for production ranking."
)
add_para(
    "(2) Stability — hyperparameters were selected under 5-fold "
    "StratifiedKFold cross-validation; the CV precision (optimized "
    "scoring) was consistent before the held-out test evaluation, "
    "confirming the model generalizes rather than overfitting."
)
add_para(
    "(3) Business alignment — the precision–recall operating point "
    "(threshold 0.294) is calibrated so that ~50% of flagged "
    "invoices are genuine disputes, doubling the hit-rate of "
    "random outreach. The Logistic Regression is retained as an "
    "explanation layer (signed coefficients) but is not the "
    "operational ranker because its precision (0.299) implies "
    "2 wasted collector calls for every real dispute found."
)

add_heading("2.4 Predictive Modeling Approaches", level=2)
add_heading("2.4.1 Human-Implemented Model", level=3)
add_para(
    "Logistic Regression was selected for interpretability — "
    "collectors need to know why an invoice was flagged. Its "
    "coefficients give a direct, signed contribution per feature "
    "and the model is robust on small, mixed-type tabular data."
)
add_heading("2.4.2 AI-Recommended Model", level=3)
add_para(
    "Random Forest was recommended by the AI tool for its ability "
    "to capture non-linear interactions (e.g., Country × Amount × "
    "CustomerFreq) without manual feature crosses, and for built-in "
    "feature-importance reporting that supports business "
    "explanation."
)

add_heading("2.5 Model Comparison and Initial Evaluation", level=2)
add_para(
    "Both models were trained on identical features. Evaluation "
    "metrics: Accuracy, Precision, Recall, F1, and ROC-AUC. "
    "Precision is the primary business metric — collectors have "
    "limited capacity, so flagged invoices must be truly disputable "
    "to avoid wasted outreach. ROC-AUC measures threshold-"
    "independent ranking quality."
)

add_heading("2.6 Model Refinement and Performance Improvement",
            level=2)
add_para(
    "Three refinements were applied: (a) SelectKBest (ANOVA F-test) "
    "to drop weak features; (b) GridSearchCV with 5-fold "
    "StratifiedKFold to tune Logistic Regression (C, k) on F1 and "
    "Random Forest (n_estimators, max_depth, min_samples_split, k) "
    "on Precision; (c) probability-threshold calibration via the "
    "precision–recall curve, with the Random Forest threshold moved "
    "from 0.50 to 0.294 to hit a target precision of 0.50 while "
    "preserving recall. The refined Random Forest is the final "
    "model — it dominates every metric except recall (where LR is "
    "higher but with very low precision) and offers the best "
    "ranking quality (ROC-AUC 0.746)."
)

# ===== 3. RESULTS =====
add_heading("3. Results", level=1)
add_para(
    "Table I summarizes the four models on the held-out test set "
    "(20% stratified). The refined Random Forest is the "
    "best-performing model on Accuracy, Precision, F1, and ROC-AUC."
)

# Results table
table = doc.add_table(rows=5, cols=6)
table.style = "Light Grid Accent 1"
hdr = table.rows[0].cells
for i, h in enumerate(["Model", "Accuracy", "Precision", "Recall",
                       "F1", "ROC-AUC"]):
    hdr[i].text = h
rows = [
    ("LR — Baseline", "0.5870", "0.2946", "0.5893", "0.3929", "0.6192"),
    ("RF — Baseline", "0.7126", "0.3707", "0.3839", "0.3772", "0.6200"),
    ("LR — Refined", "0.5931", "0.2986", "0.5893", "0.3964", "0.6151"),
    ("RF — Refined", "0.7713", "0.4946", "0.4107", "0.4488", "0.7460"),
]
for i, row in enumerate(rows, start=1):
    for j, v in enumerate(row):
        table.rows[i].cells[j].text = v
# Shrink table font
for r_ in table.rows:
    for c_ in r_.cells:
        for p_ in c_.paragraphs:
            for run_ in p_.runs:
                run_.font.size = Pt(8)
                run_.font.name = "Times New Roman"

cap = doc.add_paragraph()
cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
cap.paragraph_format.space_after = Pt(6)
cr = cap.add_run("Table I. Test-set performance of baseline vs refined models.")
cr.italic = True
cr.font.size = Pt(9)
cr.font.name = "Times New Roman"

add_figure(IMG / "image1.png",
           "Figure 1. Baseline vs Refined — metric-wise comparison.",
           width_inches=3.2)

add_para(
    "The Refined Random Forest lifts Precision from 0.371 to 0.495 "
    "(+33%) and ROC-AUC from 0.620 to 0.746 (+20%) versus its "
    "baseline. Recall dips slightly (0.384 → 0.411 after threshold "
    "calibration, still positive), which is acceptable because the "
    "collections queue is precision-bound."
)

add_figure(IMG / "image2.png",
           "Figure 2. Confusion matrices — baseline vs refined "
           "(LR top, RF bottom).",
           width_inches=3.2)

add_para(
    "Refined RF correctly catches 46 of 112 disputes while raising "
    "True-Negatives from 309 to 335 and lowering False-Positives "
    "from 73 to 47 — exactly the trade-off collectors want."
)

add_figure(IMG / "image3.png",
           "Figure 3. ROC curves — Refined RF dominates all other "
           "models.",
           width_inches=3.0)

add_figure(IMG / "image4.png",
           "Figure 4. Refined Random Forest — feature importance.",
           width_inches=3.2)

add_para(
    "InvoiceAmount, CustomerFreq, and countryCode together account "
    "for ~71% of the model's predictive signal — large invoices, "
    "less-frequent customers, and certain countries dominate "
    "dispute risk."
)

add_figure(IMG / "image5.png",
           "Figure 5. Refined Logistic Regression — standardized "
           "coefficients.",
           width_inches=3.2)

add_para(
    "LR coefficients confirm the same direction: InvoiceAmount, "
    "InvoiceMonth and countryCode push risk up; PaperlessBill and "
    "DaysSincePaperless push risk down (long-time paperless "
    "customers dispute less)."
)

# ===== 3.1 DISCUSSION =====
add_heading("3.1 Discussion", level=2)

add_para("Model Performance.", bold=True, space_before=4)
add_para(
    "The Refined Random Forest is the strongest model: ROC-AUC "
    "0.746 is well above the 0.62 baselines and the only score that "
    "exceeds the 0.70 'usable for production ranking' threshold. "
    "Precision rose from 0.371 to 0.495 because GridSearchCV "
    "optimized for precision and threshold calibration trimmed the "
    "low-confidence region of the score distribution."
)

add_para("Model Suitability.", bold=True, space_before=4)
add_para(
    "Collections capacity is the binding constraint. A "
    "precision-bound, ranked-risk queue means almost half of the "
    "flagged invoices are real disputes, doubling the hit-rate of "
    "random outreach (22.7%). This directly meets the business "
    "objective of reducing wasted collector effort while still "
    "catching ~41% of disputes early."
)

add_para("Comparison of Models.", bold=True, space_before=4)
add_para(
    "Logistic Regression remains useful for explanation but its "
    "low Precision (≤0.30) makes it unsuitable as the operational "
    "ranker. Random Forest was the right AI recommendation for "
    "this tabular, mixed-signal problem; refinement amplified its "
    "advantage on every metric except recall."
)

add_para("Feature Importance and Business Insight.", bold=True, space_before=4)
add_para(
    "Both models agree on the top drivers — InvoiceAmount, "
    "CustomerFreq, countryCode — which is reassuring (consistency "
    "across model families). Translating these into actionable "
    "business intelligence:"
)
add_para(
    "\u2022 InvoiceAmount (most important in RF, strongest positive "
    "coefficient in LR): large invoices are disproportionately "
    "disputed. Collections teams should prioritize the top-value "
    "flagged invoices first — a single avoided dispute on a "
    "high-amount invoice can offset the cost of several collector "
    "calls."
)
add_para(
    "\u2022 CustomerFreq (second in RF): low-frequency (new or "
    "infrequent) customers dispute more often. This signals that "
    "onboarding and first-invoice communication should be "
    "reinforced for new accounts to reduce early-tenure disputes."
)
add_para(
    "\u2022 countryCode (top-3 in both models): dispute risk differs "
    "significantly by geography. The collections team can create "
    "country-specific follow-up protocols for the highest-risk "
    "regions identified by the model."
)
add_para(
    "\u2022 PaperlessBill and DaysSincePaperless (negative LR "
    "coefficients): customers enrolled in paperless billing for "
    "longer dispute less. Proactively enrolling new customers in "
    "e-invoicing is a low-cost, model-supported intervention to "
    "reduce dispute rates over time."
)
add_para(
    "Overall, the model converts 22.7% random dispute prevalence "
    "into a ranked queue where ~50% of the top-flagged invoices "
    "are real disputes — a 2.2x lift that directly reduces "
    "wasted collector effort and shortens Days Sales Outstanding."
)

# ===== 3.2 Conclusion =====
add_heading("3.2 Conclusion and Recommendations", level=2)
add_para("Conclusion.", bold=True, space_before=4)
add_para(
    "A leakage-free predictive pipeline was built and refined on "
    "2,466 AR invoices. The Refined Random Forest (ROC-AUC 0.746, "
    "Precision 0.495) is the final model. It is interpretable "
    "through its feature importances, stable under 5-fold CV, and "
    "aligned with the business need for a precision-oriented "
    "early-warning queue."
)
add_para("Recommendations.", bold=True, space_before=4)
add_para(
    "(1) Deploy the Refined RF as a daily-scored ranked queue at "
    "operating-point threshold 0.294. (2) Monitor monthly for "
    "drift on InvoiceAmount distribution and customer mix. "
    "(3) Re-train quarterly as new settled invoices grow the "
    "dataset. (4) Pair the model with the LR coefficient view as "
    "the explanation layer for collectors and auditors."
)

# ===== References =====
add_heading("References", level=1)
add_para("[1] IBM, \"Sample dataset: WA_Fn-UseC_-Accounts-Receivable\", "
         "IBM Cognos Analytics samples.")
add_para("[2] F. Pedregosa et al., \"Scikit-learn: Machine Learning in "
         "Python,\" J. Mach. Learn. Res., vol. 12, pp. 2825–2830, 2011.")
add_para("[3] L. Breiman, \"Random Forests,\" Machine Learning, vol. 45, "
         "no. 1, pp. 5–32, 2001.")
add_para("[4] D. W. Hosmer, S. Lemeshow, and R. X. Sturdivant, "
         "Applied Logistic Regression, 3rd ed. Hoboken, NJ: Wiley, 2013.")
add_para("[5] J. Davis and M. Goadrich, \"The relationship between "
         "Precision–Recall and ROC curves,\" in Proc. ICML, 2006, "
         "pp. 233–240.")

doc.save(str(OUT))
print("Saved:", OUT)
print("Paragraphs:", len(doc.paragraphs), "Tables:", len(doc.tables))
